#!/usr/bin/env python3
"""
LuxeReserve - Advanced Database Technical Report & Presentation Generator
Generates:
  1. Technical Report (DOCX) - Formal academic deliverable
  2. Presentation (PPTX) - For project defense
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pptx import Presentation
from pptx.util import Inches as PInches, Pt as PPt
from pptx.dml.color import RGBColor as PRGBColor
from pptx.enum.text import PP_ALIGN

# ============================================================
# COLOR PALETTE (LuxeReserve Brand)
# ============================================================
DEEP_TEAL = RGBColor(0x14, 0x3D, 0x42)
WARM_BEIGE = RGBColor(0xF4, 0xED, 0xE3)
BRONZE = RGBColor(0xB5, 0x79, 0x3F)
TEXT_STRONG = RGBColor(0x17, 0x28, 0x2B)
TEXT_SOFT = RGBColor(0x60, 0x71, 0x72)
TEXT_ON_DARK = RGBColor(0xF7, 0xF1, 0xE7)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF0, 0xF0, 0xF0)
MEDIUM_GRAY = RGBColor(0xCC, 0xCC, 0xCC)

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def add_formatted_paragraph(doc, text, bold=False, size=12, color=None, alignment=None, space_after=6):
    """Add a paragraph with consistent formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text, level=0):
    """Add a bullet point with consistent formatting."""
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_code_block(doc, code_text):
    """Add a code block with monospace font."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_STRONG
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    return p

# ============================================================
# PART 1: TECHNICAL REPORT (DOCX)
# ============================================================

def create_technical_report():
    doc = Document()
    
    # ---- Page Setup ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    
    # ---- Styles ----
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = TEXT_STRONG
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    for level in range(1, 4):
        h_style = doc.styles['Heading %d' % level]
        h_font = h_style.font
        h_font.name = 'Times New Roman'
        h_font.color.rgb = DEEP_TEAL
        h_font.bold = True
        if level == 1:
            h_font.size = Pt(16)
            h_style.paragraph_format.space_before = Pt(24)
            h_style.paragraph_format.space_after = Pt(12)
        elif level == 2:
            h_font.size = Pt(14)
            h_style.paragraph_format.space_before = Pt(18)
            h_style.paragraph_format.space_after = Pt(8)
        else:
            h_font.size = Pt(12)
            h_style.paragraph_format.space_before = Pt(12)
            h_style.paragraph_format.space_after = Pt(6)
    
    # ============================================================
    # COVER PAGE
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ADVANCED DATABASE\nTECHNICAL REPORT')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = DEEP_TEAL
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2500' * 40)
    run.font.color.rgb = BRONZE
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('LuxeReserve Hotel Reservation Management System')
    run.font.size = Pt(18)
    run.font.color.rgb = DEEP_TEAL
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Polyglot Persistence Architecture\nSQL Server 2022 + MongoDB Atlas')
    run.font.size = Pt(14)
    run.font.color.rgb = TEXT_SOFT
    run.font.name = 'Times New Roman'
    
    for _ in range(4):
        doc.add_paragraph()
    
    info_lines = [
        'Course: Advanced Database (Co so du lieu nang cao)',
        'Group: DAF04',
        'Date: April 2026',
        'Version: 1.0',
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)
        run.font.color.rgb = TEXT_SOFT
        run.font.name = 'Times New Roman'
    
    doc.add_page_break()
    
    # ============================================================
    # TABLE OF CONTENTS
    # ============================================================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('1.', 'Introduction', 1),
        ('1.1', 'Project Overview', 2),
        ('1.2', 'System Objectives', 2),
        ('1.3', 'Scope and Limitations', 2),
        ('2.', 'Technology Stack', 1),
        ('2.1', 'Frontend Technologies', 2),
        ('2.2', 'Backend Technologies', 2),
        ('2.3', 'Database Technologies', 2),
        ('2.4', 'Polyglot Persistence Strategy', 2),
        ('3.', 'System Architecture', 1),
        ('3.1', 'Three-Tier Architecture', 2),
        ('3.2', 'Hybrid Data Layer Design', 2),
        ('3.3', 'API Design and Routing', 2),
        ('3.4', 'Authentication and Authorization', 2),
        ('4.', 'Database Design', 1),
        ('4.1', 'Entity-Relationship Model', 2),
        ('4.2', 'SQL Server Schema (30 Tables)', 2),
        ('4.3', 'MongoDB Collections (3 Collections)', 2),
        ('4.4', 'Domain Decomposition', 2),
        ('5.', 'Advanced Database Techniques', 1),
        ('5.1', 'Recursive CTE for Location Hierarchy', 2),
        ('5.2', 'Pessimistic Locking for Concurrency Control', 2),
        ('5.3', 'Optimistic Locking for Inventory Updates', 2),
        ('5.4', 'Window Functions for Revenue Analytics', 2),
        ('5.5', 'Triggers for Audit and Data Integrity', 2),
        ('5.6', 'Stored Procedures for Business Logic', 2),
        ('5.7', 'Views for Financial Reporting', 2),
        ('5.8', 'Computed Columns for Derived Data', 2),
        ('6.', 'System Features', 1),
        ('6.1', 'Reservation Management', 2),
        ('6.2', 'Payment Processing', 2),
        ('6.3', 'Front Desk Operations', 2),
        ('6.4', 'Housekeeping and Maintenance', 2),
        ('6.5', 'Guest Services and Loyalty', 2),
        ('6.6', 'Admin Dashboard and Reporting', 2),
        ('7.', 'Source Code Architecture', 1),
        ('7.1', 'Backend Structure (Node.js/Express)', 2),
        ('7.2', 'Frontend Structure (React/Vite)', 2),
        ('7.3', 'Database Migration Scripts', 2),
        ('7.4', 'Testing Framework', 2),
        ('8.', 'Conclusion', 1),
        ('8.1', 'Achievements', 2),
        ('8.2', 'Challenges and Lessons Learned', 2),
        ('8.3', 'Future Enhancements', 2),
    ]
    for num, title, level in toc_items:
        p = doc.add_paragraph()
        indent = '    ' * (level - 1)
        run = p.add_run('%s%s %s' % (indent, num, title))
        run.font.size = Pt(12 if level == 1 else 11)
        run.font.bold = (level == 1)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_page_break()
    
    # ============================================================
    # CHAPTER 1: INTRODUCTION
    # ============================================================
    doc.add_heading('1. Introduction', level=1)
    
    doc.add_heading('1.1 Project Overview', level=2)
    doc.add_paragraph(
        'LuxeReserve is a comprehensive Hotel Reservation Management System (HRMS) designed for '
        'luxury hotel chains operating across Southeast Asia. The system provides an end-to-end '
        'solution for managing hotel inventory, guest reservations, payment processing, front desk '
        'operations, housekeeping coordination, maintenance tracking, and revenue analytics. '
        'Built as an academic project for the Advanced Database course (DAF04), LuxeReserve '
        'demonstrates the practical application of advanced database techniques including recursive '
        'Common Table Expressions (CTEs), pessimistic and optimistic locking mechanisms, window '
        'functions for analytical processing, triggers for automated audit logging, stored procedures '
        'for complex business transactions, and a polyglot persistence architecture combining '
        'SQL Server 2022 with MongoDB Atlas.'
    )
    doc.add_paragraph(
        'The system serves three primary user groups: guests who book rooms through a public web '
        'portal, front desk staff who manage check-in and check-out operations and room assignments, '
        'and administrators who oversee inventory, pricing, and generate business intelligence reports. '
        'With 30 SQL Server tables, 3 MongoDB collections, 82 API endpoints, and a React-based '
        'frontend, LuxeReserve represents a production-grade implementation of a real-world '
        'hospitality management platform.'
    )
    
    doc.add_heading('1.2 System Objectives', level=2)
    objectives = [
        'Implement a robust reservation system with ACID-compliant transactions to prevent double-booking and ensure data integrity.',
        'Demonstrate advanced database techniques including recursive CTEs, window functions, triggers, stored procedures, and locking mechanisms.',
        'Design a polyglot persistence architecture that leverages SQL Server for transactional data and MongoDB for flexible content management.',
        'Provide a comprehensive set of RESTful APIs for hotel management, reservation processing, payment handling, and operational workflows.',
        'Build an intuitive user interface for guests, front desk staff, and administrators with role-based access control.',
        'Implement real-time inventory management with pessimistic locking to handle concurrent booking requests safely.',
        'Create an audit trail system that automatically logs all critical data changes through database triggers.',
    ]
    for obj in objectives:
        add_bullet(doc, obj)
    
    doc.add_heading('1.3 Scope and Limitations', level=2)
    doc.add_paragraph(
        'The scope of LuxeReserve encompasses direct booking management, room inventory control, '
        'check-in and check-out workflows, room transfers, payment processing (including VNPay integration), '
        'housekeeping task management, maintenance ticket tracking, guest service ordering, invoice '
        'generation, loyalty program management, and revenue analytics reporting. The system supports '
        'multi-hotel, multi-brand, and multi-chain operations with a hierarchical location structure.'
    )
    doc.add_paragraph(
        'The following features are outside the current scope: integration with third-party Online Travel '
        'Agencies (OTAs) such as Booking.com or Agoda, native mobile applications, full human resources '
        'management, enterprise accounting integration, artificial intelligence-based revenue management, '
        'and multi-property consolidated accounting. These limitations represent opportunities for future '
        'enhancement rather than deficiencies in the current implementation.'
    )
    
    doc.add_page_break()
    
    # ============================================================
    # CHAPTER 2: TECHNOLOGY STACK
    # ============================================================
    doc.add_heading('2. Technology Stack', level=1)
    
    doc.add_heading('2.1 Frontend Technologies', level=2)
    doc.add_paragraph(
        'The frontend of LuxeReserve is built using React 18 with Vite as the build tool, providing '
        'a modern, component-based Single Page Application (SPA) architecture. React Router v6 handles '
        'client-side navigation across 12 distinct page routes including the home page, search results, '
        'hotel details, booking flow, reservation lookup, account management, and separate portals for '
        'administrators and cashier staff. State management is handled through React Context API with '
        'two primary contexts: AuthContext for authentication state and FlashContext for global toast '
        'notifications. The user interface follows a premium hospitality design system with a warm beige '
        'background palette, deep teal primary colors, and bronze accent elements, creating an elegant '
        'and sophisticated visual experience befitting a luxury hotel brand.'
    )
    
    doc.add_heading('2.2 Backend Technologies', level=2)
    doc.add_paragraph(
        'The backend is implemented using Node.js with the Express framework, providing a RESTful API '
        'server that exposes 82 endpoints across 14 route groups. The server connects to SQL Server 2022 '
        'Express using the mssql package (with msnodesqlv8 ODBC driver for Windows Authentication support) '
        'and to MongoDB Atlas using the official MongoDB Node.js driver v6. Authentication is handled '
        'through JSON Web Tokens (JWT) with bcryptjs for password hashing at cost factor 10. The system '
        'integrates with VNPay for Vietnamese payment gateway processing and Nodemailer with Gmail SMTP '
        'for transactional email notifications including booking confirmations and password reset OTPs.'
    )
    
    doc.add_heading('2.3 Database Technologies', level=2)
    doc.add_paragraph(
        'LuxeReserve employs a polyglot persistence architecture utilizing two database systems in parallel. '
        'SQL Server 2022 Express serves as the primary transactional database, handling all ACID-critical '
        'operations including reservations, payments, inventory management, and user accounts. The SQL Server '
        'schema comprises 30 tables organized into 12 logical domains, with comprehensive foreign key '
        'relationships, CHECK constraints, computed columns, and indexes to ensure data integrity and '
        'query performance. MongoDB Atlas serves as the flexible content store, housing rich descriptive '
        'content for hotels, room types, and amenities in three collections: Hotel_Catalog, room_type_catalog, '
        'and amenity_master. The two databases are linked through business keys (amenity_code, room_type_code, '
        'hotel_id) and merged at the API layer to provide a unified data model to the frontend.'
    )
    
    doc.add_heading('2.4 Polyglot Persistence Strategy', level=2)
    doc.add_paragraph(
        'The decision to use two database systems was driven by the distinct nature of the data being managed. '
        'Transactional data such as reservations, payments, and inventory require strict ACID compliance, '
        'referential integrity, and support for advanced locking mechanisms, making SQL Server the appropriate '
        'choice. Content data such as hotel descriptions, image galleries, amenity features, and room type '
        'details benefit from MongoDB\'s flexible schema, embedded document model, and rich text search '
        'capabilities. This hybrid approach ensures that each type of data is stored in the most suitable '
        'environment, with the API layer responsible for merging results from both databases when serving '
        'composite requests such as hotel detail pages that combine operational data from SQL Server with '
        'rich content from MongoDB.'
    )
    
    doc.add_page_break()
    
    # ============================================================
    # CHAPTER 3: SYSTEM ARCHITECTURE
    # ============================================================
    doc.add_heading('3. System Architecture', level=1)
    
    doc.add_heading('3.1 Three-Tier Architecture', level=2)
    doc.add_paragraph(
        'LuxeReserve follows a classic three-tier architecture pattern. The presentation tier consists of '
        'a React SPA running on Vite dev server (port 5173) that communicates with the application tier '
        'through HTTP REST calls with JWT Bearer token authentication. The application tier is a Node.js '
        'Express server (port 3000) that implements business logic, authentication middleware, role-based '
        'access control, and data aggregation across both database systems. The data tier comprises SQL '
        'Server 2022 Express for transactional data and MongoDB Atlas for content data, with the application '
        'tier acting as the bridge between the two. This separation of concerns allows each tier to be '
        'developed, tested, and scaled independently.'
    )
    
    doc.add_heading('3.2 Hybrid Data Layer Design', level=2)
    doc.add_paragraph(
        'The hybrid data layer is a cornerstone of the LuxeReserve architecture. SQL Server handles all '
        'operational and transactional data including hotel configurations, room inventory, guest profiles, '
        'reservations, payments, invoices, housekeeping tasks, maintenance tickets, and audit logs. MongoDB '
        'stores rich content including hotel descriptions with embedded image galleries, room type features '
        'and photographs, and amenity master data with icons and tags. When a client requests hotel details, '
        'the API layer executes parallel queries to both databases, then merges the results using business '
        'key mapping. For example, a hotel detail endpoint retrieves operational data from SQL Server '
        '(Hotel, Brand, Chain, Location, RoomType, RoomRate, HotelAmenity tables) and rich content from '
        'MongoDB (Hotel_Catalog, room_type_catalog, amenity_master collections), combining them into a '
        'unified response. This approach ensures that 94 percent of API endpoints are served primarily '
        'from SQL Server, with only 6 percent requiring hybrid data merging.'
    )
    
    doc.add_heading('3.3 API Design and Routing', level=2)
    doc.add_paragraph(
        'All API endpoints are prefixed with /api/v1/ and organized into 14 route groups: hotels, rooms, '
        'guests, auth, promotions, reservations, payments, admin, locations, services, housekeeping, '
        'maintenance, invoices, and vnpay. Each route group is implemented as a separate Express router '
        'module in the src/routes/ directory. The API follows consistent response conventions: successful '
        'responses return { success: true, data: {...} } while error responses return { success: false, '
        'message: "..." } with appropriate HTTP status codes (200 for reads, 201 for creates, 400 for '
        'validation errors, 401 for authentication failures, 403 for authorization failures, 404 for '
        'not found, and 500 for server errors). Business logic is separated from route handlers and '
        'implemented in the src/services/ directory, ensuring clean separation of concerns.'
    )
    
    doc.add_heading('3.4 Authentication and Authorization', level=2)
    doc.add_paragraph(
        'Authentication is implemented using JSON Web Tokens (JWT) with an 8-hour expiration period. '
        'The system supports two user types: system users (staff members with roles such as ADMIN, CASHIER, '
        'FRONT_DESK) and guests (hotel customers). A unified login endpoint auto-detects the user type '
        'and returns the appropriate token with role information. The middleware stack consists of three '
        'layers: attachAuthContext (optional JWT decoding for public endpoints), requireAuth (mandatory '
        'authentication), and requireSystemUser (verification that the authenticated user is a staff member). '
        'Role-based access control is enforced at the endpoint level, with ADMIN users having full access, '
        'CASHIER and FRONT_DESK users having operational access, and guests having access only to their '
        'own reservations and profile data. Demo accounts include admin/admin (ADMIN), cashier/cashier '
        '(CASHIER, FRONT_DESK), and dqc/dqc (guest, PLATINUM loyalty tier).'
    )
    
    doc.add_page_break()
    
    # ============================================================
    # CHAPTER 4: DATABASE DESIGN
    # ============================================================
    doc.add_heading('4. Database Design', level=1)
    
    doc.add_heading('4.1 Entity-Relationship Model', level=2)
    doc.add_paragraph(
        'The LuxeReserve database schema comprises 30 SQL Server tables organized into 12 logical domains, '
        'each representing a distinct functional area of the hotel management system. The domains include '
        'Location Hierarchy (self-referencing Location table with 5 levels), Hotel Organization (HotelChain, '
        'Brand, Hotel with supporting HotelPolicy and HotelAmenity), Room Management (RoomType, Room, '
        'RoomFeature, RoomAvailability), Guest Management (Guest with computed full_name column, '
        'GuestAddress, GuestPreference, GuestAuth, LoyaltyAccount), System Users and Roles (SystemUser, '
        'Role, UserRole), Pricing and Promotions (RatePlan, RoomRate, RateChangeLog, Promotion, '
        'BookingChannel), Reservation Core (Reservation, ReservationRoom, ReservationGuest, '
        'ReservationStatusHistory), Financial (Payment, PaymentCardToken, Invoice), Stay and Services '
        '(StayRecord, ServiceCatalog, ReservationService, HotelReview), Operations (HousekeepingTask, '
        'MaintenanceTicket), Loyalty (LoyaltyRedemption), and Audit and Locks (InventoryLockLog, AuditLog). '
        'The complete ERD is documented in Mermaid format in the project documentation.'
    )
    
    doc.add_heading('4.2 SQL Server Schema (30 Tables)', level=2)
    doc.add_paragraph(
        'The SQL Server schema is designed with careful attention to normalization, data integrity, and '
        'performance. Each table uses BIGINT identity columns for primary keys, VARCHAR for code-based '
        'identifiers, NVARCHAR for multilingual text content, and appropriate data types for numeric and '
        'temporal data. Foreign key relationships are explicitly defined with cascading constraints where '
        'appropriate. CHECK constraints enforce domain integrity on status fields, type classifications, '
        'and business rules. Unique constraints prevent duplicate records on business key combinations. '
        'Indexes are created on frequently queried columns and foreign key columns to optimize query '
        'performance. Notable design decisions include the use of computed persisted columns for Guest.full_name, '
        'denormalized hotel_id in RoomAvailability and HousekeepingTask for query performance, snapshot '
        'currency_code fields in Reservation, Payment, and Invoice for historical accuracy, and the '
        'CHECK constraint on RoomFeature ensuring at least one foreign key reference is not null.'
    )
    
    tables_info = [
        ('Location', 'Self-referencing hierarchy for geographical regions, countries, states, cities, and districts. Supports recursive CTE queries.'),
        ('HotelChain', 'Luxury hotel chain master data (e.g., Marriott International, IHG).'),
        ('Brand', 'Hotel brand under a chain (e.g., Ritz-Carlton, W Hotels, InterContinental).'),
        ('Hotel', 'Physical hotel properties with operational details, currency, and location reference.'),
        ('HotelPolicy', 'Hotel-specific policies for cancellation, deposits, children, pets, and smoking.'),
        ('HotelAmenity', 'Mapping between hotels and amenities with pricing and availability information.'),
        ('RoomType', 'Commercial room categories (Deluxe, Suite, Villa, Presidential Suite) with capacity and bed configuration.'),
        ('Room', 'Physical room instances with status tracking for occupancy, housekeeping, and maintenance.'),
        ('RoomFeature', 'Features assigned to either a specific room or a room type, with CHECK constraint ensuring at least one reference.'),
        ('RoomAvailability', 'Daily inventory records per room supporting pessimistic locking with UPDLOCK and HOLDLOCK hints.'),
        ('RatePlan', 'Pricing strategies (BAR, Non-Refundable, Member, Package, Corporate, Promo) with meal inclusion options.'),
        ('RoomRate', 'Daily room pricing by room type and rate plan, with audit trail through RateChangeLog.'),
        ('RateChangeLog', 'Audit log for rate changes exceeding 50 percent threshold, populated by the Price Integrity Guard trigger.'),
        ('Promotion', 'Marketing promotions applicable to hotels or brands with discount values and eligibility rules.'),
        ('BookingChannel', 'Distribution channels (Direct, OTA, GDS, Corporate, Travel Agent, Wholesaler) with commission rates.'),
        ('Reservation', 'Core booking header with guest information, dates, financial totals, and status tracking.'),
        ('ReservationRoom', 'Room line items within a reservation with rate snapshots and occupancy tracking.'),
        ('ReservationGuest', 'Guest information snapshot for each reservation, supporting multiple guests per booking.'),
        ('ReservationStatusHistory', 'Complete status change history for each reservation, auto-populated by trigger.'),
        ('Guest', 'Guest profiles with computed persisted full_name column and contact information.'),
        ('GuestAddress', 'Guest addresses categorized as home, work, or billing with primary flag.'),
        ('GuestPreference', 'Guest preferences for bed type, pillow, floor, diet, view, and other personalized services.'),
        ('GuestAuth', 'Guest authentication credentials with email verification and account status tracking.'),
        ('LoyaltyAccount', 'Loyalty program membership with tier levels (Silver, Gold, Platinum, Black) and points balance.'),
        ('LoyaltyRedemption', 'Points redemption records linking guests, promotions, and reservations.'),
        ('Payment', 'Payment transactions with type classification (Deposit, Prepayment, Full Payment, Refund, Incidental Hold).'),
        ('PaymentCardToken', 'Tokenized payment card information for PCI compliance and recurring billing.'),
        ('Invoice', 'Financial invoices generated from the vw_ReservationTotal view with status workflow.'),
        ('ServiceCatalog', 'Hotel service catalog including spa, dining, airport transfer, butler, and other luxury services.'),
        ('ReservationService', 'Service orders attached to reservations with pricing and status tracking.'),
        ('StayRecord', 'Actual stay records capturing check-in and check-out times linked to reservation rooms.'),
        ('HotelReview', 'Guest reviews and ratings for hotels with moderation workflow.'),
        ('HousekeepingTask', 'Housekeeping task management with priority levels and staff assignment.'),
        ('MaintenanceTicket', 'Maintenance issue tracking with severity levels and resolution workflow.'),
        ('InventoryLockLog', 'Audit log for pessimistic locking operations during room booking.'),
        ('AuditLog', 'General audit log for tracking data changes across sensitive tables via triggers.'),
        ('SystemUser', 'Staff user accounts with department assignment and account status.'),
        ('Role', 'Role definitions for role-based access control.'),
        ('UserRole', 'User-to-role assignments for authorization management.'),
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = 'Table Name'
    hdr[1].text = 'Description'
    for name, desc in tables_info:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('4.3 MongoDB Collections (3 Collections)', level=2)
    doc.add_paragraph(
        'MongoDB Atlas hosts three collections that complement the SQL Server schema. The Hotel_Catalog '
        'collection stores rich hotel content including short and long descriptions, highlight bullet points, '
        'image galleries with captions and categories, embedded amenity information with icons and tags, '
        'embedded room type descriptions with feature flags and images, and detailed location and contact '
        'information. The room_type_catalog collection stores room type descriptions, boolean feature flags '
        '(has_balcony, has_private_pool, has_lounge_access, has_butler_service), image galleries, and '
        'highlight text. The amenity_master collection stores amenity names, descriptions, icons, tags as '
        'arrays, and illustrative images. These collections are linked to SQL Server tables through business '
        'keys: amenity_code links HotelAmenity to amenity_master, room_type_code links RoomType to '
        'room_type_catalog, and hotel_id links Hotel to Hotel_Catalog.'
    )
    
    doc.add_heading('4.4 Domain Decomposition', level=2)
    doc.add_paragraph(
        'The database schema is decomposed into 12 logical domains to manage complexity and ensure '
        'cohesive grouping of related tables. Domain 1 (Location Hierarchy) contains the self-referencing '
        'Location table. Domain 2 (Hotel Organization) contains HotelChain, Brand, Hotel, HotelPolicy, '
        'and HotelAmenity. Domain 3 (Room Management) contains RoomType, Room, RoomFeature, and '
        'RoomAvailability. Domain 4 (Guest Management) contains Guest, GuestAddress, GuestPreference, '
        'GuestAuth, and LoyaltyAccount. Domain 5 (System Users) contains SystemUser, Role, and UserRole. '
        'Domain 6 (Pricing) contains RatePlan, RoomRate, RateChangeLog, Promotion, and BookingChannel. '
        'Domain 7 (Reservation Core) contains Reservation, ReservationRoom, ReservationGuest, and '
        'ReservationStatusHistory. Domain 8 (Financial) contains Payment, PaymentCardToken, and Invoice. '
        'Domain 9 (Stay and Services) contains StayRecord, ServiceCatalog, ReservationService, and '
        'HotelReview. Domain 10 (Operations) contains HousekeepingTask and MaintenanceTicket. Domain 11 '
        '(Loyalty) contains LoyaltyRedemption. Domain 12 (Audit and Locks) contains InventoryLockLog '
        'and AuditLog. This decomposition facilitates modular development, testing, and maintenance.'
    )
    
    doc.add_page_break()
    
    # ============================================================
    # CHAPTER 5: ADVANCED DATABASE TECHNIQUES
    # ============================================================
    doc.add_heading('5. Advanced Database Techniques', level=1)
    
    doc.add_heading('5.1 Recursive CTE for Location Hierarchy', level=2)
    doc.add_paragraph(
        'The Location table implements an adjacency list model for hierarchical geographical data, '
        'supporting five levels: Region (level 0), Country (level 1), State or Province (level 2), '
        'City (level 3), and District (level 4). A recursive Common Table Expression (CTE) named '
        'vw_LocationTree traverses this hierarchy starting from any root node, computing the depth '
        'from the starting point and constructing a full path string that concatenates ancestor names. '
        'The recursive CTE consists of an anchor member that selects root locations (where '
        'parent_location_id IS NULL) and a recursive member that joins child locations to their '
        'parents. The final view includes a hotel_count aggregation that counts hotels at each '
        'location level, enabling queries such as "find all hotels in Southeast Asia" without '
        'multiple round-trips to the database. This technique is demonstrated in the API endpoint '
        'GET /api/v1/locations/tree and is documented in the vw_LocationTree view definition.'
    )
    
    doc.add_heading('5.2 Pessimistic Locking for Concurrency Control', level=2)
    doc.add_paragraph(
        'To prevent double-booking in a concurrent environment, LuxeReserve implements pessimistic '
        'locking using SQL Server\'s UPDLOCK and HOLDLOCK table hints within the sp_ReserveRoom and '
        'sp_TransferRoom stored procedures. When a booking request is received, the procedure begins '
        'a transaction and issues a SELECT statement with WITH (UPDLOCK, HOLDLOCK) on the RoomAvailability '
        'row for the requested room and date. UPDLOCK prevents other transactions from acquiring update '
        'or exclusive locks on the same row, while HOLDLOCK ensures the lock is held until the transaction '
        'commits or rolls back (equivalent to SERIALIZABLE isolation level on that row). If the room is '
        'available (status is OPEN), the procedure updates the status to BOOKED, logs the successful '
        'lock in InventoryLockLog, and commits the transaction. If another concurrent request attempts '
        'to book the same room, it is blocked at the SELECT statement until the first transaction '
        'completes, at which point it sees the updated BOOKED status and rejects the request. This '
        'mechanism guarantees that no two transactions can book the same room for the same date, '
        'eliminating race conditions entirely. The sp_TransferRoom procedure extends this concept by '
        'atomically releasing the old room and locking the new room across all nights of the stay, '
        'ensuring that room transfers are also concurrency-safe.'
    )
    
    doc.add_heading('5.3 Optimistic Locking for Inventory Updates', level=2)
    doc.add_paragraph(
        'For administrative inventory updates where contention is expected to be low, LuxeReserve '
        'employs optimistic locking using a version_no column in the RoomAvailability table. When '
        'an administrator retrieves availability data, the current version_no is included in the '
        'response. When submitting an update, the application includes the expected version_no in '
        'the UPDATE statement\'s WHERE clause: UPDATE RoomAvailability SET ... WHERE availability_id = '
        '@id AND version_no = @expected_version. If another user has modified the record in the '
        'interim, the version_no will have changed, the UPDATE will affect zero rows, and the '
        'application detects this condition and returns an HTTP 409 Conflict response. The version_no '
        'is automatically incremented (version_no = version_no + 1) on each update, ensuring that '
        'stale updates are detected and rejected. This approach provides adequate protection for '
        'low-contention administrative operations without the overhead of holding locks across '
        'transaction boundaries.'
    )
    
    doc.add_heading('5.4 Window Functions for Revenue Analytics', level=2)
    doc.add_paragraph(
        'LuxeReserve leverages SQL Server window functions extensively for revenue analytics and '
        'business intelligence reporting. The